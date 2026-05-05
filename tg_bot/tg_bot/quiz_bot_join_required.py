import asyncio
import json
import logging
import math
import os
import random
import re
import time
from contextlib import suppress
from dataclasses import asdict, dataclass, field
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Callable, Optional

from aiogram import Bot, Dispatcher, F, Router
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.exceptions import TelegramBadRequest
from aiogram.filters import Command, CommandStart
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    CallbackQuery,
    Document as TgDocument,
    FSInputFile,
    Message,
    PhotoSize,
    WebAppInfo,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder
from dotenv import load_dotenv

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except Exception:
    load_workbook = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


# =========================
# CONFIGURATION
# =========================

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")

BOT_TOKEN = os.getenv("BOT_TOKEN", "").strip()
ADMIN_IDS = {
    int(x.strip())
    for x in os.getenv("ADMIN_IDS", "").split(",")
    if x.strip().isdigit()
}

# Stickers
STICKER_WELCOME = os.getenv("STICKER_WELCOME", "").strip()
STICKER_SUCCESS = os.getenv("STICKER_SUCCESS", "").strip()
STICKER_ERROR = os.getenv("STICKER_ERROR", "").strip()
STICKER_PREMIUM = os.getenv("STICKER_PREMIUM", "").strip()
STICKER_TEST = os.getenv("STICKER_TEST", "").strip()

# Required join gate
REQUIRED_JOIN_CHAT = os.getenv("REQUIRED_JOIN_CHAT", "@talabaga_yordam_pro").strip()
REQUIRED_JOIN_URL = os.getenv("REQUIRED_JOIN_URL", "https://t.me/talabaga_yordam_pro").strip()
REQUIRED_JOIN_TITLE = os.getenv("REQUIRED_JOIN_TITLE", "Talabaga Yordam Pro").strip()
WEB_APP_URL = os.getenv("WEB_APP_URL", "https://quizgen-ai.vercel.app").strip()

# Paths
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
PAYMENTS_UPLOAD_DIR = DATA_DIR / "payment_receipts"
for folder in [DATA_DIR, UPLOAD_DIR, PAYMENTS_UPLOAD_DIR]:
    folder.mkdir(parents=True, exist_ok=True)

# Files
USERS_FILE = DATA_DIR / "users.json"
TESTS_FILE = DATA_DIR / "tests.json"
PAYMENTS_FILE = DATA_DIR / "payments.json"
ATTEMPTS_FILE = DATA_DIR / "attempts.json"

# Limits
SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".txt"}
MAX_FILE_SIZE_MB = 20
MAX_RECEIPT_SIZE_MB = 20
MAX_QUESTIONS_FREE = int(os.getenv("MAX_QUESTIONS_FREE", "50") or "50")

# Payment
PAYMENT_CARD = os.getenv("PAYMENT_CARD", "").strip()
PAYMENT_OWNER = os.getenv("PAYMENT_OWNER", "Admin").strip()

# Tariffs
TARIFFS = {
    3: {
        "days": 3,
        "title_uz": "3 kun Premium",
        "title_ru": "Премиум на 3 дня",
        "title_en": "3-day Premium",
        "base_price": 15000,
        "final_price": 15000,
        "discount_percent": 0,
    },
    5: {
        "days": 5,
        "title_uz": "5 kun Premium",
        "title_ru": "Премиум на 5 дней",
        "title_en": "5-day Premium",
        "base_price": 25000,
        "final_price": 25000,
        "discount_percent": 0,
    },
    7: {
        "days": 7,
        "title_uz": "7 kun Premium",
        "title_ru": "Премиум на 7 дней",
        "title_en": "7-day Premium",
        "base_price": 35000,
        "final_price": 30000,
        "discount_percent": 15,
    },
    14: {
        "days": 14,
        "title_uz": "14 kun Premium",
        "title_ru": "Премиум на 14 дней",
        "title_en": "14-day Premium",
        "base_price": 70000,
        "final_price": 60000,
        "discount_percent": 15,
    },
}
TARIFF_ORDER = [3, 5, 7, 14]

logger = logging.getLogger("quiz_bot")
router = Router()
SESSION_KEY = "active_session"
_timer_tasks: dict[int, asyncio.Task] = {}


# =========================
# JSON STORE
# =========================

class JsonStore:
    def __init__(self, path: Path, default_factory: Callable = dict, cache_ttl: float = 2.0):
        self._path = path
        self._default_factory = default_factory
        self._cache_ttl = cache_ttl
        self._lock = asyncio.Lock()
        self._data: Any = None
        self._last_load = 0.0

    async def _load_unlocked(self) -> Any:
        if self._path.exists():
            try:
                content = await asyncio.to_thread(self._path.read_text, encoding="utf-8")
                self._data = json.loads(content)
            except Exception:
                self._data = self._default_factory()
        else:
            self._data = self._default_factory()
        self._last_load = time.monotonic()
        return self._data

    async def _load(self) -> Any:
        now = time.monotonic()
        if self._data is not None and (now - self._last_load) < self._cache_ttl:
            return self._data
        async with self._lock:
            if self._data is not None and (time.monotonic() - self._last_load) < self._cache_ttl:
                return self._data
            return await self._load_unlocked()

    async def get(self) -> Any:
        return await self._load()

    async def update(self, mutator: Callable[[Any], Any]) -> Any:
        async with self._lock:
            data = await self._load_unlocked()
            result = mutator(data)
            await self._persist(data)
            return result

    async def _persist(self, data: Any) -> None:
        def _write() -> None:
            tmp = self._path.with_suffix(".tmp")
            tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
            tmp.replace(self._path)

        await asyncio.to_thread(_write)


_users_store = JsonStore(USERS_FILE, dict)
_tests_store = JsonStore(TESTS_FILE, dict)
_payments_store = JsonStore(PAYMENTS_FILE, dict)
_attempts_store = JsonStore(ATTEMPTS_FILE, dict)


# =========================
# I18N
# =========================

I18N = {
    "uz": {
        "choose_language": "🌐 Tilni tanlang",
        "language_saved": "✅ Til saqlandi: O'zbekcha",
        "main_menu": "🏠 Asosiy menyu",
        "upload_file": "📂 Fayl yuklash",
        "my_tests": "📚 Mening testlarim",
        "premium": "⭐ Premium",
        "help": "ℹ️ Yordam",
        "back": "⬅️ Orqaga",
        "cancel": "❌ Bekor qilish",
        "web_app": "🌐 Web Ilovada ishlash",
        "admin": "🛠 Admin panel",
        "send_file": (
            "📎 Test faylini yuboring.\n\n"
            "Qo'llab-quvvatlanadi: <b>docx / pdf / xlsx / txt</b>\n"
            "Maksimal hajm: <b>{max_mb} MB</b>"
        ),
        "bad_file": "❌ Faqat docx, pdf, xlsx yoki txt fayl yuboring.",
        "bad_receipt": "❌ Chekni rasm (jpg/png) yoki PDF ko'rinishida yuboring.",
        "file_too_large": "❌ Fayl juda katta. Maksimal: {max_mb} MB.",
        "analyzing": "⏳ Fayl tahlil qilinmoqda, biroz kuting...",
        "no_questions": (
            "❌ Savollar topilmadi.\n\n"
            "Faylda quyidagi formatlardan biri bo'lishi kerak:\n"
            "1) Tayyor A/B/C/D testlar\n"
            "2) Savol # Javob ko'rinishi\n"
            "3) Oxirida javoblar kaliti"
        ),
        "analysis_result": (
            "✅ Tahlil tugadi.\n"
            "📚 Topilgan savollar: <b>{count}</b>\n"
            "⚠️ Shubhali elementlar: <b>{suspicious}</b>"
        ),
        "ask_test_name": "📝 Test nomini kiriting:",
        "ask_chunk_size": "🧩 Har bir bo'limda nechta savol bo'lsin?\nMasalan: <b>30</b>, <b>50</b>, <b>100</b>",
        "ask_order": "🔀 Savollar tartibini tanlang:",
        "sequential": "📑 Ketma-ket",
        "random": "🎲 Random",
        "ask_timer_type": "⏱ Vaqt turini tanlang:",
        "per_question": "Har savol uchun",
        "whole_test": "Butun test uchun",
        "ask_timer_value": "⏰ Vaqtni kiriting.\nMasalan: <b>30s</b>, <b>1m</b>, <b>1.5m</b>, <b>45m</b>",
        "created": (
            "✅ Test tayyor!\n\n"
            "📌 Nomi: <b>{name}</b>\n"
            "📚 Jami savollar: <b>{count}</b>\n"
            "🧩 Bo'limlar: <b>{sections}</b>\n"
            "📦 Har bo'lim: <b>{chunk}</b>\n"
            "⏱ Rejim: <b>{timer}</b>"
        ),
        "empty_tests": "📭 Sizda hali testlar yo'q.",
        "your_tests": "📚 Sizning testlaringiz:",
        "choose_section": "📂 Bo'limni tanlang:",
        "question_header": (
            "<b>{title}</b>\n"
            "Savol {current}/{total} | ⏱ {time_left} | ✅ {correct} ❌ {wrong}\n\n"
            "{question_text}\n\n"
            "{options}"
        ),
        "finished": (
            "🏁 Test tugadi!\n\n"
            "✅ To'g'ri: <b>{correct}</b>\n"
            "❌ Noto'g'ri: <b>{wrong}</b>\n"
            "📊 Natija: <b>{percent}%</b>\n"
            "🏅 Daraja: <b>{grade}</b>"
        ),
        "free_limit": "⭐ Bepul rejimda faqat {max} tagacha savol saqlanadi. Ko'proq savol uchun Premium kerak.",
        "invalid_number": "❌ To'g'ri son kiriting.",
        "invalid_time": "❌ Vaqt formati noto'g'ri. Masalan: 30s, 1m, 1.5m, 45m",
        "time_up": "⏰ Vaqt tugadi!",
        "section_label": "{idx}-bo'lim",
        "session_expired": "⚠️ Sessiya tugagan. Testni qayta boshlang.",
        "cancelled": "❌ Bekor qilindi.",
        "not_found": "Topilmadi",
        "section_not_found": "Bo'lim topilmadi",
        "answer_correct": "✅ To'g'ri",
        "answer_wrong": "❌ Noto'g'ri. To'g'ri javob: {correct}",
        "loading": "⏳ Yuklanmoqda...",
        "admin_only": "Bu bo'lim faqat admin uchun.",
        "status": "📌 Holat",
        "reject": "❌ Rad etish",
        "view_pending": "💸 Kutayotgan to'lovlar",
        "back_menu": "🏠 Menyuga qaytish",
        "back_to_premium": "⬅️ Premiumga qaytish",
        "tariffs": "💳 Tariflar",
        "send_receipt": "📎 Chek yuborish",
        "stop_test": "⛔ Testni to'xtatish",
        "back_to_tests": "⬅️ Testlarga qaytish",
        "my_results": "📈 Natijalarim",
        "no_results_yet": "📭 Hali natijalar yo'q.",
        "help_text": (
            "🤖 Bot ishlash tartibi:\n\n"
            "1) Fayl yuklaysiz\n"
            "2) Bot savollarni ajratadi\n"
            "3) Test nomi, bo'lim hajmi va vaqtni tanlaysiz\n"
            "4) Testni Telegram ichida ishlaysiz"
        ),
        "premium_info": (
            "⭐ <b>Premium</b>\n\n"
            "Premiumda:\n"
            "• 50 tadan ko'p savol bilan ishlash\n"
            "• Kattaroq test to'plamlari\n"
            "• Kelajakda qo'shimcha imkoniyatlar\n\n"
            "Tarifni tanlang 👇"
        ),
        "payment_created": "✅ To'lov arizangiz qabul qilindi. Admin tasdiqlashini kuting.",
        "payment_already_pending": "⏳ Sizda allaqachon ko'rib chiqilayotgan premium ariza bor.",
        "payment_rejected_user": "❌ Premium arizangiz rad etildi. Iltimos, aniq va haqiqiy chek yuboring.",
        "payment_not_found": "Topilmadi yoki allaqachon ko'rib chiqilgan.",
        "premium_status_inactive": "ℹ️ Sizda faol premium yo'q.",
        "premium_status_pending": "⏳ Premium arizangiz ko'rib chiqilmoqda.",
        "admin_stats": (
            "🛠 <b>Admin panel</b>\n\n"
            "👥 Userlar: <b>{users}</b>\n"
            "📚 Testlar: <b>{tests}</b>\n"
            "💸 Kutayotgan to'lovlar: <b>{pending}</b>\n"
            "⭐ Faol premium userlar: <b>{premium_users}</b>"
        ),
        "admin_pending_empty": "📭 Kutayotgan premium arizalar yo'q.",
        "admin_pending_list": "💸 Kutayotgan premium arizalar:",
        "join_required": (
            "🔒 Botdan to'liq foydalanish uchun avval quyidagi guruhga qo'shiling:\n"
            "<a href=\"{url}\">{title}</a>\n\n"
            "Qo'shilganingizdan keyin <b>Tekshirish</b> tugmasini bosing."
        ),
        "join_group": "👥 Guruhga qo'shilish",
        "check_join": "✅ Tekshirish",
        "join_success": "✅ A'zolik tasdiqlandi.",
        "join_failed": "❌ Siz hali guruhga qo'shilmagansiz yoki bot tekshira olmadi.",
    },
    "ru": {
        "choose_language": "🌐 Выберите язык",
        "language_saved": "✅ Язык сохранён: Русский",
        "main_menu": "🏠 Главное меню",
        "upload_file": "📂 Загрузить файл",
        "my_tests": "📚 Мои тесты",
        "premium": "⭐ Премиум",
        "help": "ℹ️ Помощь",
        "back": "⬅️ Назад",
        "cancel": "❌ Отмена",
        "web_app": "🌐 Работа в Web версии",
        "admin": "🛠 Админ панель",
        "send_file": (
            "📎 Отправьте файл с тестами.\n\n"
            "Поддерживается: <b>docx / pdf / xlsx / txt</b>\n"
            "Максимальный размер: <b>{max_mb} MB</b>"
        ),
        "bad_file": "❌ Отправьте только docx, pdf, xlsx или txt.",
        "bad_receipt": "❌ Отправьте чек как фото (jpg/png) или PDF.",
        "file_too_large": "❌ Файл слишком большой. Максимум: {max_mb} MB.",
        "analyzing": "⏳ Идёт анализ файла, подождите...",
        "no_questions": (
            "❌ Вопросы не найдены.\n\n"
            "В файле должен быть один из форматов:\n"
            "1) Тесты с вариантами A/B/C/D\n"
            "2) Вопрос # ответ\n"
            "3) Ключ ответов в конце"
        ),
        "analysis_result": (
            "✅ Анализ завершён.\n"
            "📚 Найдено вопросов: <b>{count}</b>\n"
            "⚠️ Подозрительных элементов: <b>{suspicious}</b>"
        ),
        "ask_test_name": "📝 Введите название теста:",
        "ask_chunk_size": "🧩 Сколько вопросов должно быть в каждом разделе?\nНапример: <b>30</b>, <b>50</b>, <b>100</b>",
        "ask_order": "🔀 Выберите порядок вопросов:",
        "sequential": "📑 По порядку",
        "random": "🎲 Случайно",
        "ask_timer_type": "⏱ Выберите тип времени:",
        "per_question": "На каждый вопрос",
        "whole_test": "На весь тест",
        "ask_timer_value": "⏰ Введите время.\nНапример: <b>30s</b>, <b>1m</b>, <b>1.5m</b>, <b>45m</b>",
        "created": (
            "✅ Тест готов!\n\n"
            "📌 Название: <b>{name}</b>\n"
            "📚 Всего вопросов: <b>{count}</b>\n"
            "🧩 Разделов: <b>{sections}</b>\n"
            "📦 В каждом разделе: <b>{chunk}</b>\n"
            "⏱ Режим: <b>{timer}</b>"
        ),
        "empty_tests": "📭 У вас пока нет тестов.",
        "your_tests": "📚 Ваши тесты:",
        "choose_section": "📂 Выберите раздел:",
        "question_header": (
            "<b>{title}</b>\n"
            "Вопрос {current}/{total} | ⏱ {time_left} | ✅ {correct} ❌ {wrong}\n\n"
            "{question_text}\n\n"
            "{options}"
        ),
        "finished": (
            "🏁 Тест завершён!\n\n"
            "✅ Верно: <b>{correct}</b>\n"
            "❌ Неверно: <b>{wrong}</b>\n"
            "📊 Результат: <b>{percent}%</b>\n"
            "🏅 Уровень: <b>{grade}</b>"
        ),
        "free_limit": "⭐ В бесплатном режиме сохраняется не больше {max} вопросов. Для больших наборов нужен Премиум.",
        "invalid_number": "❌ Введите корректное число.",
        "invalid_time": "❌ Неверный формат времени. Например: 30s, 1m, 1.5m, 45m",
        "time_up": "⏰ Время вышло!",
        "section_label": "Раздел {idx}",
        "session_expired": "⚠️ Сессия истекла. Запустите тест снова.",
        "cancelled": "❌ Отменено.",
        "not_found": "Не найдено",
        "section_not_found": "Раздел не найден",
        "answer_correct": "✅ Верно",
        "answer_wrong": "❌ Неверно. Правильный ответ: {correct}",
        "loading": "⏳ Загрузка...",
        "admin_only": "Этот раздел только для администратора.",
        "status": "📌 Статус",
        "reject": "❌ Отклонить",
        "view_pending": "💸 Ожидающие оплаты",
        "back_menu": "🏠 В меню",
        "back_to_premium": "⬅️ Назад к премиуму",
        "tariffs": "💳 Тарифы",
        "send_receipt": "📎 Отправить чек",
        "stop_test": "⛔ Остановить тест",
        "back_to_tests": "⬅️ Назад к тестам",
        "my_results": "📈 Мои результаты",
        "no_results_yet": "📭 Результатов пока нет.",
        "help_text": (
            "🤖 Как работает бот:\n\n"
            "1) Вы загружаете файл\n"
            "2) Бот выделяет вопросы\n"
            "3) Вы выбираете название, размер раздела и время\n"
            "4) Решаете тест прямо в Telegram"
        ),
        "premium_info": (
            "⭐ <b>Премиум</b>\n\n"
            "В премиуме:\n"
            "• больше 50 вопросов\n"
            "• большие тестовые наборы\n"
            "• будущие возможности\n\n"
            "Выберите тариф 👇"
        ),
        "payment_created": "✅ Ваша заявка на оплату принята. Дождитесь проверки админом.",
        "payment_already_pending": "⏳ У вас уже есть заявка на рассмотрении.",
        "payment_rejected_user": "❌ Ваша заявка отклонена. Пожалуйста, отправляйте только понятный и настоящий чек.",
        "payment_not_found": "Не найдено или уже обработано.",
        "premium_status_inactive": "ℹ️ У вас нет активного премиума.",
        "premium_status_pending": "⏳ Ваша заявка на премиум проверяется.",
        "admin_stats": (
            "🛠 <b>Админ панель</b>\n\n"
            "👥 Пользователи: <b>{users}</b>\n"
            "📚 Тесты: <b>{tests}</b>\n"
            "💸 Ожидают оплаты: <b>{pending}</b>\n"
            "⭐ Активных premium: <b>{premium_users}</b>"
        ),
        "admin_pending_empty": "📭 Нет ожидающих заявок.",
        "admin_pending_list": "💸 Ожидающие заявки:",
        "join_required": (
            "🔒 Чтобы полностью пользоваться ботом, сначала вступите в группу:\n"
            "<a href=\"{url}\">{title}</a>\n\n"
            "После вступления нажмите кнопку <b>Проверить</b>."
        ),
        "join_group": "👥 Вступить в группу",
        "check_join": "✅ Проверить",
        "join_success": "✅ Участие подтверждено.",
        "join_failed": "❌ Вы ещё не вступили в группу или бот не смог это проверить.",
    },
    "en": {
        "choose_language": "🌐 Choose language",
        "language_saved": "✅ Language saved: English",
        "main_menu": "🏠 Main menu",
        "upload_file": "📂 Upload file",
        "my_tests": "📚 My tests",
        "premium": "⭐ Premium",
        "help": "ℹ️ Help",
        "back": "⬅️ Back",
        "cancel": "❌ Cancel",
        "web_app": "🌐 Work in Web App",
        "admin": "🛠 Admin panel",
        "send_file": (
            "📎 Send your test file.\n\n"
            "Supported: <b>docx / pdf / xlsx / txt</b>\n"
            "Maximum size: <b>{max_mb} MB</b>"
        ),
        "bad_file": "❌ Please send only docx, pdf, xlsx or txt.",
        "bad_receipt": "❌ Please send receipt as photo (jpg/png) or PDF.",
        "file_too_large": "❌ File is too large. Maximum: {max_mb} MB.",
        "analyzing": "⏳ Analyzing file, please wait...",
        "no_questions": (
            "❌ No questions found.\n\n"
            "The file should contain one of these formats:\n"
            "1) A/B/C/D tests\n"
            "2) Question # answer\n"
            "3) Answer key at the end"
        ),
        "analysis_result": (
            "✅ Analysis completed.\n"
            "📚 Questions found: <b>{count}</b>\n"
            "⚠️ Suspicious items: <b>{suspicious}</b>"
        ),
        "ask_test_name": "📝 Enter a test title:",
        "ask_chunk_size": "🧩 How many questions per section?\nExamples: <b>30</b>, <b>50</b>, <b>100</b>",
        "ask_order": "🔀 Choose question order:",
        "sequential": "📑 Sequential",
        "random": "🎲 Random",
        "ask_timer_type": "⏱ Choose timer type:",
        "per_question": "Per question",
        "whole_test": "Whole test",
        "ask_timer_value": "⏰ Enter time.\nExamples: <b>30s</b>, <b>1m</b>, <b>1.5m</b>, <b>45m</b>",
        "created": (
            "✅ Test is ready!\n\n"
            "📌 Title: <b>{name}</b>\n"
            "📚 Total questions: <b>{count}</b>\n"
            "🧩 Sections: <b>{sections}</b>\n"
            "📦 Per section: <b>{chunk}</b>\n"
            "⏱ Mode: <b>{timer}</b>"
        ),
        "empty_tests": "📭 You have no tests yet.",
        "your_tests": "📚 Your tests:",
        "choose_section": "📂 Choose a section:",
        "question_header": (
            "<b>{title}</b>\n"
            "Question {current}/{total} | ⏱ {time_left} | ✅ {correct} ❌ {wrong}\n\n"
            "{question_text}\n\n"
            "{options}"
        ),
        "finished": (
            "🏁 Test finished!\n\n"
            "✅ Correct: <b>{correct}</b>\n"
            "❌ Wrong: <b>{wrong}</b>\n"
            "📊 Result: <b>{percent}%</b>\n"
            "🏅 Grade: <b>{grade}</b>"
        ),
        "free_limit": "⭐ Free mode stores only up to {max} questions. Premium is required for larger sets.",
        "invalid_number": "❌ Enter a valid number.",
        "invalid_time": "❌ Invalid time format. Example: 30s, 1m, 1.5m, 45m",
        "time_up": "⏰ Time is up!",
        "section_label": "Section {idx}",
        "session_expired": "⚠️ Session expired. Please restart the test.",
        "cancelled": "❌ Cancelled.",
        "not_found": "Not found",
        "section_not_found": "Section not found",
        "answer_correct": "✅ Correct",
        "answer_wrong": "❌ Wrong. Correct answer: {correct}",
        "loading": "⏳ Loading...",
        "admin_only": "This section is for admins only.",
        "status": "📌 Status",
        "reject": "❌ Reject",
        "view_pending": "💸 Pending payments",
        "back_menu": "🏠 Back to menu",
        "back_to_premium": "⬅️ Back to Premium",
        "tariffs": "💳 Plans",
        "send_receipt": "📎 Send receipt",
        "stop_test": "⛔ Stop test",
        "back_to_tests": "⬅️ Back to tests",
        "my_results": "📈 My results",
        "no_results_yet": "📭 No results yet.",
        "help_text": (
            "🤖 How the bot works:\n\n"
            "1) You upload a file\n"
            "2) The bot extracts questions\n"
            "3) You choose title, section size, and time\n"
            "4) You solve the test inside Telegram"
        ),
        "premium_info": (
            "⭐ <b>Premium</b>\n\n"
            "Premium includes:\n"
            "• more than 50 questions\n"
            "• larger test sets\n"
            "• future features\n\n"
            "Choose a plan 👇"
        ),
        "payment_created": "✅ Your payment request was received. Please wait for admin confirmation.",
        "payment_already_pending": "⏳ You already have a pending payment request.",
        "payment_rejected_user": "❌ Your premium request was rejected. Please send only a real and clear receipt.",
        "payment_not_found": "Not found or already processed.",
        "premium_status_inactive": "ℹ️ You do not have an active premium plan.",
        "premium_status_pending": "⏳ Your premium request is under review.",
        "admin_stats": (
            "🛠 <b>Admin panel</b>\n\n"
            "👥 Users: <b>{users}</b>\n"
            "📚 Tests: <b>{tests}</b>\n"
            "💸 Pending payments: <b>{pending}</b>\n"
            "⭐ Active premium users: <b>{premium_users}</b>"
        ),
        "admin_pending_empty": "📭 No pending requests.",
        "admin_pending_list": "💸 Pending requests:",
        "join_required": (
            "🔒 To fully use the bot, first join this group:\n"
            "<a href=\"{url}\">{title}</a>\n\n"
            "After joining, press <b>Check</b>."
        ),
        "join_group": "👥 Join group",
        "check_join": "✅ Check",
        "join_success": "✅ Membership confirmed.",
        "join_failed": "❌ You have not joined the group yet, or the bot could not verify it.",
    },
}


async def t(user_id: int, key: str, **kwargs: Any) -> str:
    users = await _users_store.get()
    lang = users.get(str(user_id), {}).get("lang", "uz")
    text = I18N.get(lang, I18N["uz"]).get(key, key)
    try:
        return text.format(**kwargs)
    except Exception:
        return text


# =========================
# DATA MODELS
# =========================

@dataclass
class Question:
    text: str
    options: dict[str, str]
    correct: str
    explanation: str = ""

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "Question":
        return cls(
            text=data.get("text", ""),
            options=data.get("options", {}),
            correct=data.get("correct", ""),
            explanation=data.get("explanation", ""),
        )


@dataclass
class TestSet:
    test_id: str
    user_id: int
    title: str
    language: str
    original_file: str
    question_count: int
    chunk_size: int
    section_count: int
    order: str
    timer_type: str
    timer_seconds: int
    questions: list[Question] = field(default_factory=list)
    created_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class PaymentRequest:
    payment_id: str
    user_id: int
    requested_days: int
    requested_title: str
    requested_price: str
    receipt_path: str
    requested_base_price: int = 0
    requested_final_price: int = 0
    requested_discount_percent: int = 0
    receipt_kind: str = "document"
    status: str = "pending"
    approved_days: Optional[int] = None
    created_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    reviewed_at: Optional[str] = None
    reviewed_by: Optional[int] = None
    admin_message_ids: dict[str, int] = field(default_factory=dict)


class BuildTestState(StatesGroup):
    waiting_file = State()
    waiting_title = State()
    waiting_chunk_size = State()
    waiting_order = State()
    waiting_timer_type = State()
    waiting_timer_value = State()
    waiting_payment_receipt = State()
    waiting_admin_custom_days = State()


# =========================
# HELPERS
# =========================

def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def human_dt(dt_str: str) -> str:
    try:
        dt = datetime.fromisoformat(dt_str)
        return dt.astimezone().strftime("%Y-%m-%d %H:%M")
    except Exception:
        return dt_str


def format_money(amount: int) -> str:
    return f"{amount:,}".replace(",", " ") + " so'm"


def sanitize_filename(name: str) -> str:
    name = Path(name).name
    name = re.sub(r"[^a-zA-Z0-9._-]", "_", name)
    name = re.sub(r"_+", "_", name)
    return name.strip("._") or "file"


def secure_path(base_dir: Path, user_id: int, filename: str) -> Path:
    user_dir = base_dir / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    safe_name = sanitize_filename(filename)
    return user_dir / f"{int(time.time())}_{safe_name}"


def get_default_tariff() -> dict[str, Any]:
    return TARIFFS[TARIFF_ORDER[0]]


def get_tariff(days: int) -> dict[str, Any]:
    return TARIFFS.get(days, get_default_tariff())


def get_tariff_title(days: int, lang: str = "uz") -> str:
    tariff = get_tariff(days)
    if lang == "ru":
        return tariff.get("title_ru", tariff["title_uz"])
    if lang == "en":
        return tariff.get("title_en", tariff["title_uz"])
    return tariff["title_uz"]


def has_sticker_value(value: str) -> bool:
    return bool(value and value.strip())


async def send_sticker_safe(message: Message, sticker_value: str) -> None:
    if not has_sticker_value(sticker_value):
        return
    try:
        if Path(sticker_value).exists():
            await message.answer_sticker(FSInputFile(sticker_value))
        else:
            await message.answer_sticker(sticker_value)
    except Exception as e:
        logger.warning("Sticker sending failed: %s", e)


async def _safe_edit_message_text(message: Message, text: str, reply_markup: Any = None) -> None:
    try:
        await message.edit_text(text, reply_markup=reply_markup)
    except TelegramBadRequest as e:
        if "message is not modified" in str(e).lower():
            return
        try:
            await message.answer(text, reply_markup=reply_markup)
        except Exception:
            pass
    except Exception:
        try:
            await message.answer(text, reply_markup=reply_markup)
        except Exception:
            pass


def receipt_is_image(path: str | Path) -> bool:
    return Path(path).suffix.lower() in {".jpg", ".jpeg", ".png", ".webp"}


async def save_user(user_id: int, lang: Optional[str] = None) -> None:
    def mutator(users: dict[str, Any]) -> None:
        entry = users.setdefault(str(user_id), {})
        if lang:
            entry["lang"] = lang
        entry.setdefault("premium_until", None)
        entry.setdefault("premium_plan_days", None)
        entry.setdefault("premium_plan_title", None)
        entry.setdefault("created_at", now_utc().isoformat())
        entry.setdefault("join_gate_passed", None)

    await _users_store.update(mutator)


async def is_new_user(user_id: int) -> bool:
    users = await _users_store.get()
    entry = users.get(str(user_id), {})
    return not bool(entry.get("lang"))


async def get_user_lang(user_id: int) -> str:
    users = await _users_store.get()
    return users.get(str(user_id), {}).get("lang", "uz")


async def has_passed_join_gate(user_id: int) -> bool:
    users = await _users_store.get()
    value = users.get(str(user_id), {}).get("join_gate_passed", None)
    return value is not False


async def set_join_gate_passed(user_id: int, passed: bool) -> None:
    def mutator(users: dict[str, Any]) -> None:
        entry = users.setdefault(str(user_id), {})
        entry["join_gate_passed"] = passed

    await _users_store.update(mutator)


async def build_join_required_text(user_id: int) -> str:
    return await t(
        user_id,
        "join_required",
        url=REQUIRED_JOIN_URL,
        title=REQUIRED_JOIN_TITLE,
    )


async def is_user_in_required_group(bot: Bot, user_id: int) -> bool:
    try:
        member = await bot.get_chat_member(REQUIRED_JOIN_CHAT, user_id)
        status = getattr(member, "status", "")
        status = getattr(status, "value", status)
        status = str(status).lower()

        if status in {"creator", "administrator", "member"}:
            return True

        if status == "restricted" and bool(getattr(member, "is_member", False)):
            return True

        return False
    except Exception as e:
        logger.warning("Join gate check failed for user %s: %s", user_id, e)
        return False


async def ensure_join_gate_callback(callback: CallbackQuery) -> bool:
    if await has_passed_join_gate(callback.from_user.id):
        return True
    await _safe_edit_message_text(
        callback.message,
        await build_join_required_text(callback.from_user.id),
        reply_markup=await join_required_kb(callback.from_user.id),
    )
    await callback.answer()
    return False


async def ensure_join_gate_message(message: Message) -> bool:
    if await has_passed_join_gate(message.from_user.id):
        return True
    await message.answer(
        await build_join_required_text(message.from_user.id),
        reply_markup=await join_required_kb(message.from_user.id),
    )
    return False


async def is_premium(user_id: int) -> bool:
    users = await _users_store.get()
    until_str = users.get(str(user_id), {}).get("premium_until")
    if not until_str:
        return False
    try:
        until = datetime.fromisoformat(until_str)
        return until > now_utc()
    except Exception:
        return False


async def get_premium_until(user_id: int) -> Optional[str]:
    users = await _users_store.get()
    return users.get(str(user_id), {}).get("premium_until")


async def get_premium_plan_title(user_id: int) -> str:
    users = await _users_store.get()
    entry = users.get(str(user_id), {})
    lang = entry.get("lang", "uz")
    days = entry.get("premium_plan_days")
    stored_title = entry.get("premium_plan_title")
    if isinstance(days, int) and days in TARIFFS:
        return get_tariff_title(days, lang)
    if stored_title:
        return str(stored_title)
    return {"uz": "Premium", "ru": "Премиум", "en": "Premium"}.get(lang, "Premium")


async def grant_premium_days(user_id: int, days: int) -> str:
    result: dict[str, str] = {}

    def mutator(users: dict[str, Any]) -> None:
        entry = users.setdefault(str(user_id), {})
        now = now_utc()
        base = now
        if entry.get("premium_until"):
            try:
                current_until = datetime.fromisoformat(entry["premium_until"])
                if current_until > now:
                    base = current_until
            except Exception:
                pass

        new_until = base + timedelta(days=days)
        lang = entry.get("lang", "uz")
        entry["premium_until"] = new_until.isoformat()
        entry["premium_plan_days"] = days
        entry["premium_plan_title"] = get_tariff_title(days, lang)
        entry.setdefault("created_at", now.isoformat())
        result["until"] = human_dt(new_until.isoformat())

    await _users_store.update(mutator)
    return result.get("until", human_dt(now_utc().isoformat()))


async def persist_test(test_set: TestSet) -> None:
    def mutator(tests: dict[str, Any]) -> None:
        tests[test_set.test_id] = test_set.to_dict()

    await _tests_store.update(mutator)


async def get_user_tests(user_id: int) -> list[dict[str, Any]]:
    tests = await _tests_store.get()
    result = [item for item in tests.values() if item.get("user_id") == user_id]
    result.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return result


async def get_pending_payment_by_user(user_id: int) -> Optional[dict[str, Any]]:
    payments = await _payments_store.get()
    for item in payments.values():
        if item.get("user_id") == user_id and item.get("status") == "pending":
            return item
    return None


async def persist_payment(payment: PaymentRequest) -> None:
    def mutator(payments: dict[str, Any]) -> None:
        payments[payment.payment_id] = asdict(payment)

    await _payments_store.update(mutator)


def build_grade(percent: float, lang: str = "uz") -> str:
    grades = {
        "uz": [(90, "A'lo 🔥"), (75, "Yaxshi 👍"), (50, "Qoniqarli 🙂")],
        "ru": [(90, "Отлично 🔥"), (75, "Хорошо 👍"), (50, "Удовлетворительно 🙂")],
        "en": [(90, "Excellent 🔥"), (75, "Good 👍"), (50, "Satisfactory 🙂")],
    }
    for threshold, label in grades.get(lang, grades["uz"]):
        if percent >= threshold:
            return label
    return {"uz": "Yaxshilash kerak 📘", "ru": "Нужно улучшить 📘", "en": "Needs improvement 📘"}.get(lang, "Needs improvement 📘")


def parse_duration_to_seconds(raw: str) -> Optional[int]:
    raw = raw.strip().lower().replace(" ", "")
    m = re.fullmatch(r"(\d+(?:\.\d+)?)(s|m|h)", raw)
    if not m:
        return None
    value = float(m.group(1))
    if value <= 0:
        return None
    unit = m.group(2)
    if unit == "h":
        return int(value * 3600)
    if unit == "m":
        return int(value * 60)
    return int(value)


def format_seconds(seconds: int) -> str:
    minutes, sec = divmod(max(int(seconds), 0), 60)
    return f"{minutes:02d}:{sec:02d}"


def chunk_questions(questions: list[Question], chunk_size: int, order: str) -> list[list[Question]]:
    items = questions[:]
    if order == "random":
        random.shuffle(items)
    return [items[i:i + chunk_size] for i in range(0, len(items), chunk_size)]


async def build_premium_info_text(user_id: int) -> str:
    lang = await get_user_lang(user_id)
    lines = []
    for days in TARIFF_ORDER:
        tariff = get_tariff(days)
        base = int(tariff["base_price"])
        final = int(tariff["final_price"])
        discount = int(tariff.get("discount_percent", 0))
        price = f"<s>{format_money(base)}</s> → <b>{format_money(final)}</b>" if discount > 0 and final < base else f"<b>{format_money(final)}</b>"
        lines.append(f"• <b>{get_tariff_title(days, lang)}</b> — {price}")

    if lang == "ru":
        return "⭐ <b>Премиум тарифы</b>\n\n" + "\n".join(lines) + "\n\nВыберите тариф 👇"
    if lang == "en":
        return "⭐ <b>Premium plans</b>\n\n" + "\n".join(lines) + "\n\nChoose a plan 👇"
    return "⭐ <b>Premium tariflar</b>\n\n" + "\n".join(lines) + "\n\nTarifni tanlang 👇"


async def build_payment_wait_text(user_id: int, days: int) -> str:
    lang = await get_user_lang(user_id)
    tariff = get_tariff(days)
    title = get_tariff_title(days, lang)
    amount = format_money(int(tariff["final_price"]))
    if lang == "ru":
        return (
            "💳 <b>Оплата премиума</b>\n\n"
            f"📦 Тариф: <b>{title}</b>\n"
            f"💳 Карта: <code>{PAYMENT_CARD}</code>\n"
            f"👤 Получатель: <b>{PAYMENT_OWNER}</b>\n"
            f"💰 К оплате: <b>{amount}</b>\n\n"
            "После оплаты отправьте чек / скриншот."
        )
    if lang == "en":
        return (
            "💳 <b>Premium payment</b>\n\n"
            f"📦 Plan: <b>{title}</b>\n"
            f"💳 Card: <code>{PAYMENT_CARD}</code>\n"
            f"👤 Recipient: <b>{PAYMENT_OWNER}</b>\n"
            f"💰 To pay: <b>{amount}</b>\n\n"
            "After payment, send receipt / screenshot."
        )
    return (
        "💳 <b>Premium to'lovi</b>\n\n"
        f"📦 Tarif: <b>{title}</b>\n"
        f"💳 Karta: <code>{PAYMENT_CARD}</code>\n"
        f"👤 Qabul qiluvchi: <b>{PAYMENT_OWNER}</b>\n"
        f"💰 To'lanadi: <b>{amount}</b>\n\n"
        "To'lov qilgach, chek / screenshot yuboring."
    )


async def build_premium_status_text(user_id: int) -> str:
    lang = await get_user_lang(user_id)
    if await is_premium(user_id):
        until = human_dt((await get_premium_until(user_id)) or "")
        plan_title = await get_premium_plan_title(user_id)
        if lang == "ru":
            return f"✅ У вас активен премиум.\n📦 Тариф: <b>{plan_title}</b>\n⏳ Действует до: <b>{until}</b>"
        if lang == "en":
            return f"✅ Premium is active.\n📦 Plan: <b>{plan_title}</b>\n⏳ Valid until: <b>{until}</b>"
        return f"✅ Sizda premium faol.\n📦 Tarif: <b>{plan_title}</b>\n⏳ Amal qiladi: <b>{until}</b>"

    pending = await get_pending_payment_by_user(user_id)
    if pending:
        return await t(user_id, "premium_status_pending")
    return await t(user_id, "premium_status_inactive")


# =========================
# KEYBOARDS
# =========================

def language_kb() -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text="O'zbek 🇺🇿", callback_data="lang:uz")
    kb.button(text="Русский 🇷🇺", callback_data="lang:ru")
    kb.button(text="English 🇬🇧", callback_data="lang:en")
    kb.adjust(1)
    return kb.as_markup()


async def join_required_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "join_group"), url=REQUIRED_JOIN_URL)
    kb.button(text=await t(user_id, "check_join"), callback_data="join:check")
    kb.adjust(1)
    return kb.as_markup()


async def main_menu_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "upload_file"), callback_data="menu:upload")
    kb.button(text=await t(user_id, "web_app"), web_app=WebAppInfo(url=WEB_APP_URL))
    kb.button(text=await t(user_id, "my_tests"), callback_data="menu:tests")
    kb.button(text=await t(user_id, "my_results"), callback_data="menu:results")
    kb.button(text=await t(user_id, "premium"), callback_data="menu:premium")
    kb.button(text=await t(user_id, "help"), callback_data="menu:help")
    if user_id in ADMIN_IDS:
        kb.button(text=await t(user_id, "admin"), callback_data="admin:home")
    kb.adjust(1, 1, 2, 2, 1, 1)
    return kb.as_markup()


async def order_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "sequential"), callback_data="order:sequential")
    kb.button(text=await t(user_id, "random"), callback_data="order:random")
    kb.button(text=await t(user_id, "back_menu"), callback_data="menu:home")
    kb.adjust(1)
    return kb.as_markup()


async def timer_type_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "per_question"), callback_data="timer_type:per_question")
    kb.button(text=await t(user_id, "whole_test"), callback_data="timer_type:whole_test")
    kb.button(text=await t(user_id, "back_menu"), callback_data="menu:home")
    kb.adjust(1)
    return kb.as_markup()


async def section_kb(user_id: int, test_id: str, section_count: int) -> Any:
    kb = InlineKeyboardBuilder()
    for idx in range(section_count):
        kb.button(text=await t(user_id, "section_label", idx=idx + 1), callback_data=f"section:{test_id}:{idx}")
    kb.button(text=await t(user_id, "back_to_tests"), callback_data="menu:tests")
    kb.adjust(3, 1)
    return kb.as_markup()


def options_kb(test_id: str) -> Any:
    kb = InlineKeyboardBuilder()
    for letter in ["A", "B", "C", "D"]:
        kb.button(text=letter, callback_data=f"answer:{letter}")
    kb.button(text="⛔ Stop", callback_data=f"test:stop:{test_id}")
    kb.adjust(2, 2, 1)
    return kb.as_markup()


async def upload_wait_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "back"), callback_data="upload:back")
    kb.adjust(1)
    return kb.as_markup()


async def premium_home_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "status"), callback_data="premium:status")
    kb.button(text=await t(user_id, "tariffs"), callback_data="premium:tariffs")
    kb.button(text=await t(user_id, "back"), callback_data="menu:home")
    kb.adjust(1)
    return kb.as_markup()


async def premium_tariffs_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    for days in TARIFF_ORDER:
        tariff = get_tariff(days)
        kb.button(
            text=f"🗓 {days} kun — {format_money(int(tariff['final_price']))}",
            callback_data=f"premium:buy:{days}",
        )
    kb.button(text=await t(user_id, "back_to_premium"), callback_data="premium:home")
    kb.adjust(1)
    return kb.as_markup()


async def payment_wait_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "send_receipt"), callback_data="noop:receipt")
    kb.button(text=await t(user_id, "back_to_premium"), callback_data="premium:tariffs")
    kb.adjust(1)
    return kb.as_markup()


async def finished_test_kb(user_id: int, test_id: str) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "back_to_tests"), callback_data=f"open_test:{test_id}")
    kb.button(text=await t(user_id, "back_menu"), callback_data="menu:home")
    kb.adjust(1)
    return kb.as_markup()


async def admin_main_kb(user_id: int) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text=await t(user_id, "view_pending"), callback_data="admin:pending:0")
    kb.button(text=await t(user_id, "back_menu"), callback_data="menu:home")
    kb.adjust(1)
    return kb.as_markup()


async def admin_payment_actions_kb(user_id: int, payment_id: str) -> Any:
    kb = InlineKeyboardBuilder()
    kb.button(text="✅ Tasdiqlash", callback_data=f"adminpay:confirm:{payment_id}")
    kb.button(text="➕ Boshqa kun", callback_data=f"adminpay:custom:{payment_id}")
    kb.button(text=await t(user_id, "reject"), callback_data=f"adminpay:reject:{payment_id}")
    kb.button(text=await t(user_id, "admin"), callback_data="admin:home")
    kb.adjust(1)
    return kb.as_markup()


async def admin_choose_days_kb(user_id: int, payment_id: str) -> Any:
    kb = InlineKeyboardBuilder()
    for days in TARIFF_ORDER:
        kb.button(text=f"✅ {days} kun premium", callback_data=f"adminpaydays:{payment_id}:{days}")
    kb.button(text="➕ Boshqa kun", callback_data=f"adminpay:custom:{payment_id}")
    kb.button(text=await t(user_id, "reject"), callback_data=f"adminpay:reject:{payment_id}")
    kb.adjust(1)
    return kb.as_markup()


async def pending_list_kb(user_id: int, pending_items: list[dict[str, Any]], page: int, page_size: int = 10) -> Any:
    kb = InlineKeyboardBuilder()
    start = page * page_size
    end = start + page_size
    page_items = pending_items[start:end]

    for item in page_items:
        payment_id = item["payment_id"]
        short_uid = str(item["user_id"])[-4:]
        title = item.get("requested_title", "?")
        date_str = human_dt(item.get("created_at", ""))
        kb.button(text=f"👤 ...{short_uid} | {title} | {date_str}", callback_data=f"admin:payment:{payment_id}:{page}")

    if page > 0:
        kb.button(text="⬅️ Oldingi", callback_data=f"admin:pending:{page - 1}")
    if end < len(pending_items):
        kb.button(text="Keyingi ➡️", callback_data=f"admin:pending:{page + 1}")

    kb.button(text=await t(user_id, "admin"), callback_data="admin:home")
    kb.adjust(1)
    return kb.as_markup()


# =========================
# PARSERS
# =========================

QUESTION_START_RE = re.compile(r"^\s*(?:\d{1,4}|[IVXLCDM]{1,8}|(?:savol|question|вопрос)\s*\d{1,4})\s*[).:\-]\s*", re.I)
OPTION_RE = re.compile(r"^\s*([A-DА-Гa-dа-г1-4])\s*[).:\-]\s*(.+)$")
ANSWER_RE = re.compile(r"^\s*(?:javob|ответ|answer|to'g'ri\s+javob|правильный\s+ответ)\s*[:\-]\s*([A-DА-Гa-dа-г1-4])\s*$", re.I)
PIPE_RE = re.compile(r"^\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*([A-DА-Гa-dа-г1-4])\s*$")

def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\r", "\n")
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _normalize_letter(letter: str) -> str:
    mapping = {
        "A": "A", "B": "B", "C": "C", "D": "D",
        "А": "A", "Б": "B", "В": "C", "Г": "D",
        "1": "A", "2": "B", "3": "C", "4": "D",
    }
    return mapping.get(letter.strip().upper(), "")


def extract_text_from_docx(path: Path) -> str:
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(path)
    except Exception:
        return ""

    lines: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            parts = []
            for cell in row.cells:
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if (p.text or "").strip()).strip()
                if cell_text:
                    parts.append(cell_text)
            if parts:
                lines.append(" | ".join(parts))

    return "\n".join(lines)


def extract_text_from_pdf(path: Path) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(str(path))
    except Exception:
        return ""
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            pass
    return "\n".join(out)


def extract_text_from_xlsx(path: Path) -> str:
    if load_workbook is None:
        return ""
    try:
        wb = load_workbook(path, data_only=True)
    except Exception:
        return ""
    rows = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                rows.append(" | ".join(values))
    return "\n".join(rows)


def extract_text_from_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def dedupe_questions(questions: list[Question]) -> list[Question]:
    seen = set()
    out = []
    for q in questions:
        key = normalize_text(q.text).lower()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(q)
    return out


def parse_pipe_questions(text: str) -> list[Question]:
    questions = []
    for line in normalize_text(text).splitlines():
        m = PIPE_RE.match(line)
        if not m:
            continue
        correct = _normalize_letter(m.group(6))
        options = {"A": m.group(2).strip(), "B": m.group(3).strip(), "C": m.group(4).strip(), "D": m.group(5).strip()}
        if correct and all(options.values()):
            questions.append(Question(text=m.group(1).strip(), options=options, correct=correct))
    return questions


def parse_numbered_blocks(text: str) -> list[Question]:
    lines = [line.strip() for line in normalize_text(text).splitlines() if line.strip()]
    questions = []
    current_question = []
    options: dict[str, str] = {}
    correct = ""

    def flush() -> None:
        nonlocal current_question, options, correct, questions
        if current_question and len(options) == 4 and correct in {"A", "B", "C", "D"}:
            questions.append(Question(text=" ".join(current_question).strip(), options=options.copy(), correct=correct))
        current_question = []
        options = {}
        correct = ""

    for line in lines:
        if QUESTION_START_RE.match(line) and current_question and options:
            flush()

        if QUESTION_START_RE.match(line):
            if current_question and len(options) == 4:
                flush()
            current_question = [QUESTION_START_RE.sub("", line).strip()]
            continue

        opt = OPTION_RE.match(line)
        if opt:
            label = _normalize_letter(opt.group(1))
            if label:
                options[label] = opt.group(2).strip()
            continue

        ans = ANSWER_RE.match(line)
        if ans:
            correct = _normalize_letter(ans.group(1))
            continue

        if current_question and not options:
            current_question.append(line)
        elif options:
            # continuation of last option
            last_key = next(reversed(options), None)
            if last_key:
                options[last_key] = f"{options[last_key]} {line}".strip()

    flush()
    return questions


def parse_answer_key_blocks(text: str) -> list[Question]:
    normalized = normalize_text(text)
    key_map: dict[str, str] = {}
    for line in normalized.splitlines():
        for num, ans in re.findall(r"(\d+)\s*[).:\-]?\s*([A-DА-Гa-dа-г1-4])\b", line, re.I):
            key_map[num] = _normalize_letter(ans)

    if not key_map:
        return []

    pattern = re.compile(
        r"(?:^|\n)\s*(\d+)\s*[).:\-]\s*(.*?)\n"
        r"\s*[AА1]\s*[).:\-]\s*(.*?)\n"
        r"\s*[BБ2]\s*[).:\-]\s*(.*?)\n"
        r"\s*[CВ3]\s*[).:\-]\s*(.*?)\n"
        r"\s*[DГ4]\s*[).:\-]\s*(.*?)(?=\n\s*\d+\s*[).:\-]|\Z)",
        re.S | re.I,
    )
    questions = []
    for m in pattern.finditer(normalized):
        number = m.group(1)
        correct = key_map.get(number)
        if not correct:
            continue
        options = {
            "A": m.group(3).strip(),
            "B": m.group(4).strip(),
            "C": m.group(5).strip(),
            "D": m.group(6).strip(),
        }
        if all(options.values()):
            questions.append(Question(text=m.group(2).strip(), options=options, correct=correct))
    return questions


def parse_hash_pairs(text: str) -> list[tuple[str, str]]:
    pairs = []
    for line in normalize_text(text).splitlines():
        if "#" not in line:
            continue
        parts = re.split(r"\s*#{1,3}\s*", line, maxsplit=1)
        if len(parts) != 2:
            continue
        q, a = parts[0].strip(), parts[1].strip()
        if q and a:
            pairs.append((q, a))
    deduped = []
    seen = set()
    for q, a in pairs:
        key = (q.lower(), a.lower())
        if key in seen:
            continue
        seen.add(key)
        deduped.append((q, a))
    return deduped


def build_mcq_from_pairs(pairs: list[tuple[str, str]]) -> list[Question]:
    answer_pool = []
    for _, ans in pairs:
        if ans and ans.lower() not in [x.lower() for x in answer_pool]:
            answer_pool.append(ans)

    questions = []
    for q_text, correct_answer in pairs:
        distractors = [x for x in answer_pool if x.lower() != correct_answer.lower()]
        random.shuffle(distractors)
        distractors = distractors[:3]
        if len(distractors) < 3:
            continue
        options_list = distractors + [correct_answer]
        random.shuffle(options_list)
        options = {letter: options_list[i] for i, letter in enumerate(["A", "B", "C", "D"])}
        correct = next(k for k, v in options.items() if v == correct_answer)
        questions.append(Question(text=q_text, options=options, correct=correct))
    return questions


def parse_questions_from_text(text: str) -> tuple[list[Question], int]:
    suspicious = 0
    questions: list[Question] = []

    for parser in [parse_pipe_questions, parse_numbered_blocks, parse_answer_key_blocks]:
        try:
            questions.extend(parser(text))
        except Exception:
            suspicious += 1

    try:
        pairs = parse_hash_pairs(text)
        if pairs:
            questions.extend(build_mcq_from_pairs(pairs))
    except Exception:
        suspicious += 1

    questions = dedupe_questions(questions)
    return questions, suspicious


def parse_file(path: Path) -> tuple[list[Question], int]:
    ext = path.suffix.lower()
    if ext == ".docx":
        raw_text = extract_text_from_docx(path)
    elif ext == ".pdf":
        raw_text = extract_text_from_pdf(path)
    elif ext == ".xlsx":
        raw_text = extract_text_from_xlsx(path)
    elif ext == ".txt":
        raw_text = extract_text_from_txt(path)
    else:
        return [], 0
    return parse_questions_from_text(raw_text)


# =========================
# ATTEMPTS & RESULTS
# =========================

async def save_attempt_record(user_id: int, session: dict[str, Any], timeout: bool) -> None:
    total = len(session.get("questions", []))
    if total <= 0:
        return
    correct = int(session.get("correct_count", 0))
    wrong = int(session.get("wrong_count", 0))
    percent = round((correct / total) * 100, 1) if total else 0.0
    tests = await _tests_store.get()
    raw_test = tests.get(session.get("test_id"), {})
    title = raw_test.get("title", "Test")

    def mutator(attempts: dict[str, Any]) -> None:
        attempt_id = f"{user_id}_{session.get('test_id','test')}_{int(time.time())}_{random.randint(1000,9999)}"
        attempts[attempt_id] = {
            "attempt_id": attempt_id,
            "user_id": user_id,
            "test_id": session.get("test_id"),
            "test_title": title,
            "correct": correct,
            "wrong": wrong,
            "total": total,
            "percent": percent,
            "timeout": bool(timeout),
            "finished_at": now_utc().isoformat(),
        }

    await _attempts_store.update(mutator)


async def build_results_overview_text(user_id: int) -> str:
    attempts = await _attempts_store.get()
    user_attempts = [a for a in attempts.values() if int(a.get("user_id", 0)) == user_id]
    if not user_attempts:
        return await t(user_id, "no_results_yet")
    user_attempts.sort(key=lambda x: x.get("finished_at", ""), reverse=True)
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in user_attempts:
        grouped.setdefault(str(item.get("test_id")), []).append(item)
    header = await t(user_id, "my_results")
    lines = [f"📈 <b>{header}</b>"]
    for _, items in list(grouped.items())[:10]:
        last = items[0]
        best = max(items, key=lambda x: float(x.get("percent", 0)))
        lines.append(
            f"\n• <b>{last.get('test_title', 'Test')}</b>\n"
            f"  Oxirgi: <b>{last.get('percent', 0)}%</b> | Eng yaxshi: <b>{best.get('percent', 0)}%</b> | Urinish: <b>{len(items)}</b>"
        )
    return "\n".join(lines)


# =========================
# TIMER SYSTEM
# =========================

async def cancel_timer_task(user_id: int) -> None:
    task = _timer_tasks.pop(user_id, None)
    if task and not task.done():
        task.cancel()
        with suppress(asyncio.CancelledError):
            await task


def _cleanup_timer_task(user_id: int, task: asyncio.Task) -> None:
    if user_id in _timer_tasks and _timer_tasks[user_id] is task:
        _timer_tasks.pop(user_id, None)


def _make_session_dict(test_id: str, section_index: int, questions: list[Question], timer_type: str, timer_seconds: int) -> dict[str, Any]:
    return {
        "test_id": test_id,
        "section_index": section_index,
        "questions": [asdict(q) for q in questions],
        "current_index": 0,
        "correct_count": 0,
        "wrong_count": 0,
        "timer_type": timer_type,
        "timer_seconds": timer_seconds,
        "deadline_ts": time.monotonic() + timer_seconds,
        "chat_id": None,
        "message_id": None,
        "stopped": False,
    }


async def start_timer_task(bot: Bot, user_id: int, state: FSMContext) -> None:
    await cancel_timer_task(user_id)

    async def _worker() -> None:
        last_edit = 0.0
        try:
            while True:
                await asyncio.sleep(1)
                data = await state.get_data()
                session = data.get(SESSION_KEY)
                if not session or session.get("stopped"):
                    return

                now = time.monotonic()
                left = int(max(session["deadline_ts"] - now, 0))
                if left <= 0:
                    if session.get("timer_type") == "per_question":
                        still_active = await handle_question_timeout_via_bot(bot, user_id, state)
                        if not still_active:
                            return
                        last_edit = 0.0
                        continue
                    await finish_test_via_bot(bot, user_id, state, timeout=True)
                    return

                if now - last_edit >= 3:
                    await edit_active_question_via_bot(bot, user_id, state)
                    last_edit = now
        except asyncio.CancelledError:
            pass
        except Exception as exc:
            logger.warning("Timer task error for user %s: %s", user_id, exc)

    task = asyncio.create_task(_worker())
    _timer_tasks[user_id] = task
    task.add_done_callback(lambda t, uid=user_id: _cleanup_timer_task(uid, t))


async def _build_question_text(user_id: int, title: str, question: Question, current: int, total: int, time_left: str, correct: int, wrong: int) -> str:
    opts = "\n".join(f"<b>{k})</b> {v}" for k, v in question.options.items())
    return await t(
        user_id,
        "question_header",
        title=title,
        current=current,
        total=total,
        time_left=time_left,
        correct=correct,
        wrong=wrong,
        question_text=question.text,
        options=opts,
    )


async def edit_active_question_via_bot(bot: Bot, user_id: int, state: FSMContext) -> bool:
    data = await state.get_data()
    session = data.get(SESSION_KEY)
    if not session or session.get("stopped"):
        return False

    chat_id = session.get("chat_id")
    message_id = session.get("message_id")
    if not chat_id or not message_id:
        return False

    if session["current_index"] >= len(session["questions"]):
        return False

    tests = await _tests_store.get()
    title = tests.get(session["test_id"], {}).get("title", "Test")
    question = Question.from_dict(session["questions"][session["current_index"]])
    left = int(max(session["deadline_ts"] - time.monotonic(), 0))

    text = await _build_question_text(
        user_id=user_id,
        title=title,
        question=question,
        current=session["current_index"] + 1,
        total=len(session["questions"]),
        time_left=format_seconds(left),
        correct=session["correct_count"],
        wrong=session["wrong_count"],
    )

    try:
        await bot.edit_message_text(
            chat_id=chat_id,
            message_id=message_id,
            text=text,
            reply_markup=options_kb(session["test_id"]),
        )
    except TelegramBadRequest as e:
        if "message is not modified" not in str(e).lower():
            logger.warning("Timer edit error for user %s: %s", user_id, e)
    except Exception as e:
        logger.warning("Timer edit failed for user %s: %s", user_id, e)
    return True


async def handle_question_timeout_via_bot(bot: Bot, user_id: int, state: FSMContext) -> bool:
    data = await state.get_data()
    session = data.get(SESSION_KEY)
    if not session or session.get("stopped"):
        return False

    total = len(session["questions"])
    if session["current_index"] >= total:
        await finish_test_via_bot(bot, user_id, state, timeout=False)
        return False

    session["wrong_count"] += 1
    session["current_index"] += 1

    if session["current_index"] >= total:
        await state.update_data({SESSION_KEY: session})
        await finish_test_via_bot(bot, user_id, state, timeout=False)
        return False

    session["deadline_ts"] = time.monotonic() + session["timer_seconds"]
    await state.update_data({SESSION_KEY: session})
    await edit_active_question_via_bot(bot, user_id, state)
    return True


async def build_finish_text(user_id: int, session: dict[str, Any], timeout: bool) -> tuple[str, str]:
    total = len(session["questions"])
    correct = session["correct_count"]
    wrong = session["wrong_count"]
    percent = round((correct / total) * 100) if total > 0 else 0
    lang = await get_user_lang(user_id)
    grade = build_grade(percent, lang)
    test_id = session["test_id"]
    prefix = (await t(user_id, "time_up") + "\n\n") if timeout else ""
    body = await t(user_id, "finished", correct=correct, wrong=wrong, percent=percent, grade=grade)
    return prefix + body, test_id


async def finish_test_via_bot(bot: Bot, user_id: int, state: FSMContext, timeout: bool = False) -> None:
    data = await state.get_data()
    session = data.get(SESSION_KEY)
    if not session:
        return

    if timeout and session.get("timer_type") == "whole_test":
        total = len(session.get("questions", []))
        answered = int(session.get("correct_count", 0)) + int(session.get("wrong_count", 0))
        if answered < total:
            session["wrong_count"] = int(session.get("wrong_count", 0)) + (total - answered)

    await save_attempt_record(user_id, session, timeout)
    await state.update_data({SESSION_KEY: None})
    await cancel_timer_task(user_id)

    text, test_id = await build_finish_text(user_id, session, timeout)
    kb = await finished_test_kb(user_id, test_id)

    chat_id = session.get("chat_id")
    message_id = session.get("message_id")

    if chat_id and message_id:
        try:
            await bot.edit_message_text(chat_id=chat_id, message_id=message_id, text=text, reply_markup=kb)
            return
        except TelegramBadRequest as e:
            if "message is not modified" in str(e).lower():
                return
        except Exception as exc:
            logger.warning("finish_test_via_bot edit failed: %s", exc)

    if chat_id:
        try:
            await bot.send_message(chat_id=chat_id, text=text, reply_markup=kb)
        except Exception as exc:
            logger.warning("finish_test_via_bot send failed: %s", exc)


async def send_current_question(target: Message, user_id: int, session: dict[str, Any], state: FSMContext, bot: Bot, edit: bool = True) -> None:
    now_ts = time.monotonic()
    if now_ts >= session["deadline_ts"]:
        await state.update_data({SESSION_KEY: session})
        if session.get("timer_type") == "per_question":
            await handle_question_timeout_via_bot(bot, user_id, state)
        else:
            await finish_test_via_bot(bot, user_id, state, timeout=True)
        return

    tests = await _tests_store.get()
    title = tests.get(session["test_id"], {}).get("title", "Test")
    question = Question.from_dict(session["questions"][session["current_index"]])
    left = int(max(session["deadline_ts"] - now_ts, 0))
    text = await _build_question_text(
        user_id=user_id,
        title=title,
        question=question,
        current=session["current_index"] + 1,
        total=len(session["questions"]),
        time_left=format_seconds(left),
        correct=session["correct_count"],
        wrong=session["wrong_count"],
    )

    chat_id = target.chat.id
    message_id = target.message_id

    if edit:
        try:
            await target.edit_text(text, reply_markup=options_kb(session["test_id"]))
        except TelegramBadRequest as e:
            if "message is not modified" not in str(e).lower():
                sent = await target.answer(text, reply_markup=options_kb(session["test_id"]))
                chat_id = sent.chat.id
                message_id = sent.message_id
        except Exception:
            sent = await target.answer(text, reply_markup=options_kb(session["test_id"]))
            chat_id = sent.chat.id
            message_id = sent.message_id
    else:
        sent = await target.answer(text, reply_markup=options_kb(session["test_id"]))
        chat_id = sent.chat.id
        message_id = sent.message_id

    session["chat_id"] = chat_id
    session["message_id"] = message_id
    await state.update_data({SESSION_KEY: session})
    await start_timer_task(bot, user_id, state)


# =========================
# PAYMENTS / ADMIN HELPERS
# =========================

def build_admin_payment_text(payment_id: str, item: dict[str, Any]) -> str:
    return (
        f"💸 <b>Premium ariza</b>\n\n"
        f"🆔 Payment ID: <code>{payment_id}</code>\n"
        f"👤 User: <code>{item['user_id']}</code>\n"
        f"📦 So'ralgan tarif: <b>{item['requested_title']}</b>\n"
        f"💰 Narx: <b>{item['requested_price']}</b>\n"
        f"🕒 Sana: <b>{human_dt(item['created_at'])}</b>\n"
        f"📄 Holat: <b>{item['status']}</b>"
    )


def build_admin_choose_days_text(payment_id: str, item: dict[str, Any]) -> str:
    return (
        f"💳 <b>Arizani tasdiqlash</b>\n\n"
        f"🆔 Payment ID: <code>{payment_id}</code>\n"
        f"👤 User: <code>{item['user_id']}</code>\n"
        f"📦 User tanlagan tarif: <b>{item['requested_title']}</b>\n"
        f"💰 User yuborgan summa: <b>{item['requested_price']}</b>\n\n"
        "Quyidan premium muddatini tanlang yoki boshqa kun kiriting."
    )


async def send_receipt_to_admin(bot: Bot, admin_id: int, payment_id: str, item: dict[str, Any], with_actions: bool = True) -> Optional[int]:
    receipt = FSInputFile(str(item["receipt_path"]))
    caption = build_admin_payment_text(payment_id, item)
    reply_markup = await admin_payment_actions_kb(admin_id, payment_id) if with_actions and item.get("status") == "pending" else None
    try:
        if receipt_is_image(item["receipt_path"]):
            sent = await bot.send_photo(chat_id=admin_id, photo=receipt, caption=caption, reply_markup=reply_markup)
        else:
            sent = await bot.send_document(chat_id=admin_id, document=receipt, caption=caption, reply_markup=reply_markup)
        return sent.message_id
    except Exception as e:
        logger.warning("Failed to send receipt to admin %s: %s", admin_id, e)
        return None


async def cleanup_admin_receipts(bot: Bot, item: dict[str, Any], keep_chat_id: int) -> None:
    admin_message_ids: dict[str, Any] = item.get("admin_message_ids", {})
    for admin_id_str, msg_id in admin_message_ids.items():
        admin_id = int(admin_id_str)
        if admin_id == keep_chat_id:
            continue
        try:
            await bot.delete_message(chat_id=admin_id, message_id=int(msg_id))
        except Exception as exc:
            logger.warning("cleanup_admin_receipts delete failed: %s", exc)


def cleanup_receipt_file(path_str: str) -> None:
    try:
        path = Path(path_str)
        if path.exists():
            path.unlink()
    except Exception as exc:
        logger.warning("cleanup_receipt_file failed: %s", exc)


# =========================
# COMMANDS & ROUTES
# =========================

@router.message(CommandStart())
async def cmd_start(message: Message, state: FSMContext) -> None:
    await state.clear()
    await save_user(message.from_user.id)

    if await is_new_user(message.from_user.id):
        await send_sticker_safe(message, STICKER_WELCOME)
        await message.answer(await t(message.from_user.id, "choose_language"), reply_markup=language_kb())
        return

    if not await has_passed_join_gate(message.from_user.id):
        await send_sticker_safe(message, STICKER_WELCOME)
        await message.answer(await build_join_required_text(message.from_user.id), reply_markup=await join_required_kb(message.from_user.id))
        return

    await send_sticker_safe(message, STICKER_WELCOME)
    await message.answer(await t(message.from_user.id, "main_menu"), reply_markup=await main_menu_kb(message.from_user.id))


@router.message(Command("cancel"))
async def cmd_cancel(message: Message, state: FSMContext) -> None:
    await state.clear()
    await cancel_timer_task(message.from_user.id)
    await message.answer(await t(message.from_user.id, "cancelled"), reply_markup=await main_menu_kb(message.from_user.id))


@router.message(Command("admin"))
async def admin_command(message: Message) -> None:
    await save_user(message.from_user.id)
    if message.from_user.id not in ADMIN_IDS:
        await message.answer(await t(message.from_user.id, "admin_only"))
        return
    await show_admin_home(message)


@router.callback_query(F.data.startswith("lang:"))
async def set_language(callback: CallbackQuery) -> None:
    user_id = callback.from_user.id
    lang = callback.data.split(":", 1)[1]
    was_new = await is_new_user(user_id)

    await save_user(user_id, lang=lang)

    if was_new:
        await set_join_gate_passed(user_id, False)
        await _safe_edit_message_text(
            callback.message,
            await build_join_required_text(user_id),
            reply_markup=await join_required_kb(user_id),
        )
        await callback.answer(await t(user_id, "language_saved"))
        return

    await _safe_edit_message_text(callback.message, await t(user_id, "language_saved"))
    await callback.message.answer(await t(user_id, "main_menu"), reply_markup=await main_menu_kb(user_id))
    await callback.answer()


@router.callback_query(F.data == "join:check")
async def join_check(callback: CallbackQuery, bot: Bot) -> None:
    user_id = callback.from_user.id
    ok = await is_user_in_required_group(bot, user_id)
    if not ok:
        await _safe_edit_message_text(
            callback.message,
            await build_join_required_text(user_id),
            reply_markup=await join_required_kb(user_id),
        )
        await callback.answer(await t(user_id, "join_failed"), show_alert=True)
        return

    await set_join_gate_passed(user_id, True)
    await _safe_edit_message_text(
        callback.message,
        await t(user_id, "main_menu"),
        reply_markup=await main_menu_kb(user_id),
    )
    await callback.answer(await t(user_id, "join_success"))


@router.callback_query(F.data == "menu:home")
async def menu_home(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.clear()
    await cancel_timer_task(callback.from_user.id)
    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "main_menu"), reply_markup=await main_menu_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "upload:back")
async def upload_back(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.clear()
    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "main_menu"), reply_markup=await main_menu_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "premium:home")
async def premium_home(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.clear()
    await _safe_edit_message_text(callback.message, await build_premium_info_text(callback.from_user.id), reply_markup=await premium_home_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "premium:tariffs")
async def premium_tariffs(callback: CallbackQuery) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await _safe_edit_message_text(callback.message, await build_premium_info_text(callback.from_user.id), reply_markup=await premium_tariffs_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "noop:receipt")
async def noop_receipt(callback: CallbackQuery) -> None:
    await callback.answer()


@router.callback_query(F.data.startswith("test:stop:"))
async def stop_test(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.update_data({SESSION_KEY: None})
    await cancel_timer_task(callback.from_user.id)
    test_id = callback.data.split(":", 2)[2]
    test = (await _tests_store.get()).get(test_id)
    if not test:
        await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "main_menu"), reply_markup=await main_menu_kb(callback.from_user.id))
    else:
        await _safe_edit_message_text(
            callback.message,
            "⛔ Test to'xtatildi.",
            reply_markup=await section_kb(callback.from_user.id, test_id, int(test["section_count"])),
        )
    await callback.answer()


@router.callback_query(F.data == "menu:upload")
async def menu_upload(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.clear()
    await state.set_state(BuildTestState.waiting_file)
    await send_sticker_safe(callback.message, STICKER_TEST)
    await _safe_edit_message_text(
        callback.message,
        await t(callback.from_user.id, "send_file", max_mb=MAX_FILE_SIZE_MB),
        reply_markup=await upload_wait_kb(callback.from_user.id),
    )
    await callback.answer()


@router.callback_query(F.data == "menu:tests")
async def menu_tests(callback: CallbackQuery) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    user_tests = await get_user_tests(callback.from_user.id)
    if not user_tests:
        await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "empty_tests"), reply_markup=await main_menu_kb(callback.from_user.id))
        await callback.answer()
        return

    kb = InlineKeyboardBuilder()
    for item in user_tests[:25]:
        kb.button(text=f"{item['title']} ({item['question_count']})", callback_data=f"open_test:{item['test_id']}")
    kb.button(text=await t(callback.from_user.id, "back_menu"), callback_data="menu:home")
    kb.adjust(1)
    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "your_tests"), reply_markup=kb.as_markup())
    await callback.answer()


@router.callback_query(F.data == "menu:premium")
async def menu_premium(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.clear()
    await send_sticker_safe(callback.message, STICKER_PREMIUM)
    await _safe_edit_message_text(callback.message, await build_premium_info_text(callback.from_user.id), reply_markup=await premium_home_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "premium:status")
async def premium_status(callback: CallbackQuery) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    text = await build_premium_status_text(callback.from_user.id)
    await _safe_edit_message_text(callback.message, text, reply_markup=await premium_home_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data.startswith("premium:buy:"))
async def premium_buy(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    days = int(callback.data.split(":")[-1])
    existing = await get_pending_payment_by_user(callback.from_user.id)
    if existing:
        await callback.answer(await t(callback.from_user.id, "payment_already_pending"), show_alert=True)
        return

    await state.clear()
    await state.set_state(BuildTestState.waiting_payment_receipt)
    await state.update_data(requested_days=days)

    await _safe_edit_message_text(callback.message, await build_payment_wait_text(callback.from_user.id, days), reply_markup=await payment_wait_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "menu:help")
async def menu_help(callback: CallbackQuery) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "help_text"), reply_markup=await main_menu_kb(callback.from_user.id))
    await callback.answer()


@router.callback_query(F.data == "menu:results")
async def menu_results(callback: CallbackQuery) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await _safe_edit_message_text(callback.message, await build_results_overview_text(callback.from_user.id), reply_markup=await main_menu_kb(callback.from_user.id))
    await callback.answer()


@router.message(BuildTestState.waiting_file, F.document)
async def handle_file(message: Message, state: FSMContext, bot: Bot) -> None:
    if not await ensure_join_gate_message(message):
        return

    document: TgDocument = message.document
    file_name = document.file_name or "uploaded_file"
    ext = Path(file_name).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        await message.answer(await t(message.from_user.id, "bad_file"))
        return

    max_bytes = MAX_FILE_SIZE_MB * 1024 * 1024
    if document.file_size and document.file_size > max_bytes:
        await message.answer(await t(message.from_user.id, "file_too_large", max_mb=MAX_FILE_SIZE_MB))
        return

    await message.answer(await t(message.from_user.id, "analyzing"))
    path = secure_path(UPLOAD_DIR, message.from_user.id, file_name)
    await bot.download(document, destination=path)

    questions, suspicious = parse_file(path)
    if not questions:
        await send_sticker_safe(message, STICKER_ERROR)
        await message.answer(await t(message.from_user.id, "no_questions"), reply_markup=await upload_wait_kb(message.from_user.id))
        return

    if not await is_premium(message.from_user.id) and len(questions) > MAX_QUESTIONS_FREE:
        questions = questions[:MAX_QUESTIONS_FREE]
        await message.answer(await t(message.from_user.id, "free_limit", max=MAX_QUESTIONS_FREE))

    await state.update_data(
        uploaded_file=str(path),
        question_count=len(questions),
        suspicious=suspicious,
        questions=[asdict(q) for q in questions],
    )
    await state.set_state(BuildTestState.waiting_title)

    await send_sticker_safe(message, STICKER_SUCCESS)
    await message.answer(await t(message.from_user.id, "analysis_result", count=len(questions), suspicious=suspicious))
    await message.answer(await t(message.from_user.id, "ask_test_name"))


@router.message(BuildTestState.waiting_file, F.photo)
async def handle_file_photo(message: Message) -> None:
    await message.answer(await t(message.from_user.id, "bad_file"))


@router.message(BuildTestState.waiting_file)
async def handle_wrong_file(message: Message) -> None:
    await message.answer(await t(message.from_user.id, "bad_file"))


@router.message(BuildTestState.waiting_title)
async def handle_title(message: Message, state: FSMContext) -> None:
    if not await ensure_join_gate_message(message):
        return
    title = (message.text or "").strip()
    if len(title) < 2:
        await message.answer(await t(message.from_user.id, "ask_test_name"))
        return

    await state.update_data(title=title)
    await state.set_state(BuildTestState.waiting_chunk_size)
    await message.answer(await t(message.from_user.id, "ask_chunk_size"))


@router.message(BuildTestState.waiting_chunk_size)
async def handle_chunk_size(message: Message, state: FSMContext) -> None:
    if not await ensure_join_gate_message(message):
        return
    raw = (message.text or "").strip()
    if not raw.isdigit():
        await message.answer(await t(message.from_user.id, "invalid_number"))
        return

    chunk_size = int(raw)
    data = await state.get_data()
    question_count = int(data["question_count"])

    if chunk_size <= 0 or chunk_size > question_count:
        await message.answer(await t(message.from_user.id, "invalid_number"))
        return

    await state.update_data(chunk_size=chunk_size)
    await state.set_state(BuildTestState.waiting_order)
    await message.answer(await t(message.from_user.id, "ask_order"), reply_markup=await order_kb(message.from_user.id))


@router.callback_query(F.data.startswith("order:"))
async def handle_order(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.update_data(order=callback.data.split(":", 1)[1])
    await state.set_state(BuildTestState.waiting_timer_type)
    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "ask_timer_type"), reply_markup=await timer_type_kb(callback.from_user.id))
    await callback.answer()


@router.message(BuildTestState.waiting_order)
async def waiting_order_text(message: Message) -> None:
    await message.answer(await t(message.from_user.id, "ask_order"), reply_markup=await order_kb(message.from_user.id))


@router.message(BuildTestState.waiting_timer_type)
async def waiting_timer_type_text(message: Message) -> None:
    await message.answer(await t(message.from_user.id, "ask_timer_type"), reply_markup=await timer_type_kb(message.from_user.id))


@router.callback_query(F.data.startswith("timer_type:"))
async def handle_timer_type(callback: CallbackQuery, state: FSMContext) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    await state.update_data(timer_type=callback.data.split(":", 1)[1])
    await state.set_state(BuildTestState.waiting_timer_value)
    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "ask_timer_value"))
    await callback.answer()


@router.message(BuildTestState.waiting_timer_value)
async def handle_timer_value(message: Message, state: FSMContext) -> None:
    if not await ensure_join_gate_message(message):
        return
    timer_seconds = parse_duration_to_seconds(message.text or "")
    if timer_seconds is None:
        await message.answer(await t(message.from_user.id, "invalid_time"))
        return

    data = await state.get_data()
    questions = [Question.from_dict(q) for q in data["questions"]]
    title = data["title"]
    chunk_size = int(data["chunk_size"])
    order = data["order"]
    timer_type = data["timer_type"]

    test_id = f"{message.from_user.id}_{int(time.time())}"
    sections = math.ceil(len(questions) / chunk_size)

    test_set = TestSet(
        test_id=test_id,
        user_id=message.from_user.id,
        title=title,
        language=await get_user_lang(message.from_user.id),
        original_file=data["uploaded_file"],
        question_count=len(questions),
        chunk_size=chunk_size,
        section_count=sections,
        order=order,
        timer_type=timer_type,
        timer_seconds=timer_seconds,
        questions=questions,
    )

    await persist_test(test_set)
    await state.clear()
    await send_sticker_safe(message, STICKER_SUCCESS)
    await message.answer(
        await t(
            message.from_user.id,
            "created",
            name=title,
            count=len(questions),
            sections=sections,
            chunk=chunk_size,
            timer=f"{timer_type} / {format_seconds(timer_seconds)}",
        ),
        reply_markup=await section_kb(message.from_user.id, test_id, sections),
    )


@router.callback_query(F.data.startswith("open_test:"))
async def open_test(callback: CallbackQuery) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    test_id = callback.data.split(":", 1)[1]
    test = (await _tests_store.get()).get(test_id)
    if not test:
        await callback.answer(await t(callback.from_user.id, "not_found"), show_alert=True)
        return

    await _safe_edit_message_text(
        callback.message,
        await t(callback.from_user.id, "choose_section"),
        reply_markup=await section_kb(callback.from_user.id, test_id, int(test["section_count"])),
    )
    await callback.answer()


@router.callback_query(F.data.startswith("section:"))
async def start_section(callback: CallbackQuery, state: FSMContext, bot: Bot) -> None:
    if not await ensure_join_gate_callback(callback):
        return
    _, test_id, section_index_str = callback.data.split(":")
    section_index = int(section_index_str)

    raw_test = (await _tests_store.get()).get(test_id)
    if not raw_test:
        await callback.answer(await t(callback.from_user.id, "not_found"), show_alert=True)
        return

    sections = chunk_questions([Question.from_dict(q) for q in raw_test["questions"]], int(raw_test["chunk_size"]), raw_test["order"])
    if section_index >= len(sections):
        await callback.answer(await t(callback.from_user.id, "section_not_found"), show_alert=True)
        return

    session = _make_session_dict(test_id, section_index, sections[section_index], raw_test["timer_type"], int(raw_test["timer_seconds"]))
    await state.update_data({SESSION_KEY: session})
    await send_current_question(callback.message, callback.from_user.id, session, state, bot=bot, edit=True)
    await callback.answer()


@router.callback_query(F.data.startswith("answer:"))
async def answer_question(callback: CallbackQuery, state: FSMContext, bot: Bot) -> None:
    session = (await state.get_data()).get(SESSION_KEY)
    if not session:
        await callback.answer(await t(callback.from_user.id, "session_expired"), show_alert=True)
        return

    if time.monotonic() >= session["deadline_ts"]:
        await state.update_data({SESSION_KEY: session})
        if session.get("timer_type") == "per_question":
            await handle_question_timeout_via_bot(bot, callback.from_user.id, state)
        else:
            await finish_test_via_bot(bot, callback.from_user.id, state, timeout=True)
        await callback.answer()
        return

    selected = callback.data.split(":", 1)[1]
    question = Question.from_dict(session["questions"][session["current_index"]])

    if selected == question.correct:
        session["correct_count"] += 1
        await callback.answer(await t(callback.from_user.id, "answer_correct"))
    else:
        session["wrong_count"] += 1
        correct_text = question.options.get(question.correct, question.correct)
        await callback.answer(await t(callback.from_user.id, "answer_wrong", correct=f"{question.correct}) {correct_text}"), show_alert=True)

    session["current_index"] += 1
    if session["current_index"] >= len(session["questions"]):
        await state.update_data({SESSION_KEY: session})
        await finish_test_via_bot(bot, callback.from_user.id, state, timeout=False)
        return

    if session["timer_type"] == "per_question":
        session["deadline_ts"] = time.monotonic() + session["timer_seconds"]

    await state.update_data({SESSION_KEY: session})
    await send_current_question(callback.message, callback.from_user.id, session, state, bot=bot, edit=True)


# =========================
# PAYMENT FLOW
# =========================

@router.message(BuildTestState.waiting_payment_receipt, F.photo)
async def payment_receipt_photo(message: Message, state: FSMContext, bot: Bot) -> None:
    largest: PhotoSize = message.photo[-1]
    if largest.file_size and largest.file_size > MAX_RECEIPT_SIZE_MB * 1024 * 1024:
        await message.answer(await t(message.from_user.id, "file_too_large", max_mb=MAX_RECEIPT_SIZE_MB))
        return
    file_path = secure_path(PAYMENTS_UPLOAD_DIR, message.from_user.id, "receipt.jpg")
    await bot.download(largest, destination=file_path)
    await _create_payment_request_from_receipt(message, state, file_path, bot)


@router.message(BuildTestState.waiting_payment_receipt, F.document)
async def payment_receipt_document(message: Message, state: FSMContext, bot: Bot) -> None:
    document: TgDocument = message.document
    if document.file_size and document.file_size > MAX_RECEIPT_SIZE_MB * 1024 * 1024:
        await message.answer(await t(message.from_user.id, "file_too_large", max_mb=MAX_RECEIPT_SIZE_MB))
        return
    ext = Path(document.file_name or "receipt.bin").suffix.lower() or ".bin"
    file_path = secure_path(PAYMENTS_UPLOAD_DIR, message.from_user.id, f"receipt{ext}")
    await bot.download(document, destination=file_path)
    await _create_payment_request_from_receipt(message, state, file_path, bot)


@router.message(BuildTestState.waiting_payment_receipt)
async def payment_receipt_wrong(message: Message) -> None:
    await message.answer(await t(message.from_user.id, "bad_receipt"))


async def _create_payment_request_from_receipt(message: Message, state: FSMContext, file_path: Path, bot: Bot) -> None:
    existing = await get_pending_payment_by_user(message.from_user.id)
    if existing:
        cleanup_receipt_file(str(file_path))
        await message.answer(await t(message.from_user.id, "payment_already_pending"))
        return

    data = await state.get_data()
    requested_days = int(data.get("requested_days", TARIFF_ORDER[0]))
    user_lang = await get_user_lang(message.from_user.id)
    tariff = get_tariff(requested_days)
    payment_id = f"pay_{message.from_user.id}_{int(time.time())}"

    payment = PaymentRequest(
        payment_id=payment_id,
        user_id=message.from_user.id,
        requested_days=requested_days,
        requested_title=get_tariff_title(requested_days, user_lang),
        requested_price=format_money(int(tariff["final_price"])),
        receipt_path=str(file_path),
        requested_base_price=int(tariff["base_price"]),
        requested_final_price=int(tariff["final_price"]),
        requested_discount_percent=int(tariff.get("discount_percent", 0)),
        receipt_kind="photo" if receipt_is_image(file_path) else "document",
    )
    await persist_payment(payment)
    await state.clear()

    await send_sticker_safe(message, STICKER_SUCCESS)
    await message.answer(await t(message.from_user.id, "payment_created"), reply_markup=await main_menu_kb(message.from_user.id))

    admin_message_ids: dict[str, int] = {}
    for admin_id in ADMIN_IDS:
        try:
            msg_id = await send_receipt_to_admin(bot=bot, admin_id=admin_id, payment_id=payment_id, item=asdict(payment), with_actions=True)
            if msg_id:
                admin_message_ids[str(admin_id)] = int(msg_id)
        except Exception as e:
            logger.warning("Failed to notify admin %s: %s", admin_id, e)

    if admin_message_ids:
        def mutator(payments: dict[str, Any]) -> None:
            stored = payments.get(payment_id)
            if stored:
                stored["admin_message_ids"] = admin_message_ids
        await _payments_store.update(mutator)


# =========================
# ADMIN FLOW
# =========================

async def show_admin_home(target: Message | CallbackQuery) -> None:
    user_id = target.from_user.id
    users = await _users_store.get()
    tests = await _tests_store.get()
    payments = await _payments_store.get()

    pending = sum(1 for p in payments.values() if p.get("status") == "pending")
    premium_users = 0
    for _, entry in users.items():
        try:
            until_str = entry.get("premium_until")
            if until_str and datetime.fromisoformat(until_str) > now_utc():
                premium_users += 1
        except Exception:
            continue

    text = await t(user_id, "admin_stats", users=len(users), tests=len(tests), pending=pending, premium_users=premium_users)

    if isinstance(target, CallbackQuery):
        await _safe_edit_message_text(target.message, text, reply_markup=await admin_main_kb(user_id))
    else:
        await target.answer(text, reply_markup=await admin_main_kb(user_id))


@router.callback_query(F.data == "admin:home")
async def admin_home(callback: CallbackQuery) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return
    await show_admin_home(callback)
    await callback.answer()


@router.callback_query(F.data.startswith("admin:pending:"))
async def admin_pending(callback: CallbackQuery) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return

    page = int(callback.data.split(":")[-1])
    payments = await _payments_store.get()
    pending_items = [p for p in payments.values() if p.get("status") == "pending"]
    pending_items.sort(key=lambda x: x.get("created_at", ""), reverse=True)

    if not pending_items:
        await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "admin_pending_empty"), reply_markup=await admin_main_kb(callback.from_user.id))
        await callback.answer()
        return

    await _safe_edit_message_text(callback.message, await t(callback.from_user.id, "admin_pending_list"), reply_markup=await pending_list_kb(callback.from_user.id, pending_items, page))
    await callback.answer()


@router.callback_query(F.data.startswith("admin:payment:"))
async def admin_open_payment(callback: CallbackQuery) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return

    parts = callback.data.split(":")
    if len(parts) != 4:
        await callback.answer("Callback xato formatda", show_alert=True)
        return

    _, _, payment_id, _ = parts
    payments = await _payments_store.get()
    item = payments.get(payment_id)
    if not item or item.get("status") != "pending":
        await callback.answer(await t(callback.from_user.id, "payment_not_found"), show_alert=True)
        return

    await send_receipt_to_admin(bot=callback.bot, admin_id=callback.from_user.id, payment_id=payment_id, item=item, with_actions=True)
    await callback.answer()


@router.callback_query(F.data.startswith("adminpay:confirm:"))
async def admin_confirm_payment(callback: CallbackQuery) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return

    payment_id = callback.data.split(":", 2)[2]
    payments = await _payments_store.get()
    item = payments.get(payment_id)
    if not item or item.get("status") != "pending":
        await callback.answer(await t(callback.from_user.id, "payment_not_found"), show_alert=True)
        return

    await _safe_edit_message_text(callback.message, build_admin_choose_days_text(payment_id, item), reply_markup=await admin_choose_days_kb(callback.from_user.id, payment_id))
    await callback.answer("Muddatni tanlang")


@router.callback_query(F.data.startswith("adminpay:custom:"))
async def admin_custom_days_prompt(callback: CallbackQuery, state: FSMContext) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return

    payment_id = callback.data.split(":", 2)[2]
    payments = await _payments_store.get()
    item = payments.get(payment_id)
    if not item or item.get("status") != "pending":
        await callback.answer(await t(callback.from_user.id, "payment_not_found"), show_alert=True)
        return

    await state.set_state(BuildTestState.waiting_admin_custom_days)
    await state.update_data(admin_custom_payment_id=payment_id)
    text = (
        f"➕ <b>Qo'shimcha premium kun kiritish</b>\n\n"
        f"Payment ID: <code>{payment_id}</code>\n"
        f"User: <code>{item['user_id']}</code>\n\n"
        "Necha kun premium berishni yozing.\nMasalan: <b>10</b> yoki <b>21</b>"
    )
    await _safe_edit_message_text(callback.message, text, reply_markup=None)
    await callback.answer("Kun sonini yozing")


async def _approve_payment_and_notify(callback: CallbackQuery | Any, bot: Bot, payment_id: str, days: int) -> None:
    payments = await _payments_store.get()
    item = payments.get(payment_id)
    if not item or item.get("status") != "pending":
        await callback.answer(await t(callback.from_user.id, "payment_not_found"), show_alert=True)
        return

    until_text = await grant_premium_days(item["user_id"], days)
    item["status"] = "approved"
    item["approved_days"] = days
    item["reviewed_at"] = now_utc().isoformat()
    item["reviewed_by"] = callback.from_user.id

    def mutator(payments: dict[str, Any]) -> None:
        payments[payment_id] = item

    await _payments_store.update(mutator)

    user_lang = await get_user_lang(item["user_id"])
    tariff_title = get_tariff_title(days, user_lang) if days in TARIFFS else {"uz": f"{days} kun Premium", "ru": f"Премиум на {days} дней", "en": f"Premium for {days} days"}[user_lang]
    approved_text = {
        "uz": f"✅ To'lovingiz tasdiqlandi.\n📦 Tarif: <b>{tariff_title}</b>\n⏳ Amal qiladi: <b>{until_text}</b>",
        "ru": f"✅ Ваш платёж подтверждён.\n📦 Тариф: <b>{tariff_title}</b>\n⏳ Действует до: <b>{until_text}</b>",
        "en": f"✅ Your payment has been approved.\n📦 Plan: <b>{tariff_title}</b>\n⏳ Valid until: <b>{until_text}</b>",
    }[user_lang]

    try:
        await bot.send_message(chat_id=item["user_id"], text=approved_text, reply_markup=await main_menu_kb(item["user_id"]))
    except Exception as e:
        logger.warning("Failed to notify user about approved payment: %s", e)

    await cleanup_admin_receipts(bot, item, keep_chat_id=callback.from_user.id)

    done_text = (
        f"✅ Ariza tasdiqlandi.\n\n"
        f"👤 User: <code>{item['user_id']}</code>\n"
        f"📦 Berilgan premium: <b>{days} kun</b>\n"
        f"⏳ Tugash sanasi: <b>{until_text}</b>"
    )
    await _safe_edit_message_text(callback.message, done_text, reply_markup=await admin_main_kb(callback.from_user.id))
    cleanup_receipt_file(item.get("receipt_path", ""))
    await callback.answer("Tasdiqlandi ✅")


@router.callback_query(F.data.startswith("adminpaydays:"))
async def admin_approve_payment_days(callback: CallbackQuery, bot: Bot, state: FSMContext) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return

    parts = callback.data.split(":")
    if len(parts) != 3:
        await callback.answer("Callback xato formatda", show_alert=True)
        return

    _, payment_id, days_str = parts
    if not days_str.isdigit():
        await callback.answer("Kun noto'g'ri", show_alert=True)
        return

    await state.clear()
    await _approve_payment_and_notify(callback, bot, payment_id, int(days_str))


@router.message(BuildTestState.waiting_admin_custom_days)
async def admin_enter_custom_days(message: Message, state: FSMContext, bot: Bot) -> None:
    if message.from_user.id not in ADMIN_IDS:
        await state.clear()
        return

    raw = (message.text or "").strip()
    if not raw.isdigit() or int(raw) <= 0:
        await message.answer("❌ Necha kun premium berilishini son bilan yozing. Masalan: 10")
        return

    data = await state.get_data()
    payment_id = data.get("admin_custom_payment_id")
    if not payment_id:
        await state.clear()
        await message.answer("❌ Payment topilmadi.")
        return

    await state.clear()

    class _FakeCallback:
        def __init__(self, msg: Message) -> None:
            self.from_user = msg.from_user
            self.message = msg

        async def answer(self, text: str = "", show_alert: bool = False) -> None:
            return None

    fake = _FakeCallback(message)
    await _approve_payment_and_notify(fake, bot, payment_id, int(raw))


@router.callback_query(F.data.startswith("adminpay:reject:"))
async def admin_reject_payment(callback: CallbackQuery, bot: Bot, state: FSMContext) -> None:
    if callback.from_user.id not in ADMIN_IDS:
        await callback.answer(await t(callback.from_user.id, "admin_only"), show_alert=True)
        return

    payment_id = callback.data.split(":", 2)[2]
    payments = await _payments_store.get()
    item = payments.get(payment_id)
    if not item or item.get("status") != "pending":
        await callback.answer(await t(callback.from_user.id, "payment_not_found"), show_alert=True)
        return

    item["status"] = "rejected"
    item["reviewed_at"] = now_utc().isoformat()
    item["reviewed_by"] = callback.from_user.id

    def mutator(payments: dict[str, Any]) -> None:
        payments[payment_id] = item

    await _payments_store.update(mutator)
    await state.clear()

    try:
        await bot.send_message(chat_id=item["user_id"], text=await t(item["user_id"], "payment_rejected_user"), reply_markup=await main_menu_kb(item["user_id"]))
    except Exception as e:
        logger.warning("Failed to notify user about rejected payment: %s", e)

    await cleanup_admin_receipts(bot, item, keep_chat_id=callback.from_user.id)
    await _safe_edit_message_text(callback.message, "❌ Ariza rad etildi.", reply_markup=await admin_main_kb(callback.from_user.id))
    cleanup_receipt_file(item.get("receipt_path", ""))
    await callback.answer("Rad etildi ❌")

# =========================
# LOGGING / MAIN
# =========================

def configure_logging() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s")


async def on_startup(bot: Bot) -> None:
    me = await bot.get_me()
    logger.info("Bot ishga tushdi: @%s (id=%s)", me.username, me.id)
    logger.info("Adminlar: %s", sorted(ADMIN_IDS) if ADMIN_IDS else "topilmadi")
    logger.info("Majburiy group check: %s", REQUIRED_JOIN_CHAT)
    logger.info("Data papka: %s", DATA_DIR)


async def on_shutdown(bot: Bot) -> None:
    for user_id in list(_timer_tasks.keys()):
        await cancel_timer_task(user_id)
    me = await bot.get_me()
    logger.info("Bot to'xtadi: @%s", me.username)


async def main() -> None:
    configure_logging()

    if not BOT_TOKEN:
        raise RuntimeError("❌ BOT_TOKEN topilmadi. .env faylga yozing.")

    bot = Bot(token=BOT_TOKEN, default=DefaultBotProperties(parse_mode=ParseMode.HTML))
    dp = Dispatcher(storage=MemoryStorage())
    dp.include_router(router)

    try:
        await on_startup(bot)
        await dp.start_polling(bot, polling_timeout=30)
    finally:
        await on_shutdown(bot)


if __name__ == "__main__":
    asyncio.run(main())
